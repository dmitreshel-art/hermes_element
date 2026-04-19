from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import json
from datetime import datetime

json_path = '/root/element-hermes-assistant/docker-gateway/reports/acceptance_local_20260419_051026.json'
out_path = '/root/element-hermes-assistant/docker-gateway/reports/element-helpdesk-local-acceptance-20260419.docx'

with open(json_path, 'r', encoding='utf-8') as f:
    data = json.load(f)
results = data['results']
quality_checks = data['quality_checks']

summary = {'PASS': 0, 'PARTIAL': 0, 'FAIL': 0}
sections = {}
for r in results:
    summary[r['verdict']] = summary.get(r['verdict'], 0) + 1
    sec = r['section']
    if sec not in sections:
        sections[sec] = {'PASS': 0, 'PARTIAL': 0, 'FAIL': 0}
    sections[sec][r['verdict']] = sections[sec].get(r['verdict'], 0) + 1

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2)
sec.bottom_margin = Cm(2)
sec.left_margin = Cm(2)
sec.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
for name in ['Heading 1', 'Heading 2', 'Heading 3']:
    s = doc.styles[name]
    s.font.name = 'Arial'
    s.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ОТЧЁТ\nПО ЛОКАЛЬНОЙ ПРИЁМКЕ\nHERMES-ПОМОЩНИКА ПО ELEMENT')
r.bold = True
r.font.name = 'Arial'
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Проект: Element Helpdesk / matrix.22brn.ru')
r.font.name = 'Arial'
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Дата отчёта: 2026-04-19')
r.font.name = 'Arial'
r.font.size = Pt(12)


doc.add_page_break()

doc.add_heading('1. Краткое резюме', level=1)
doc.add_paragraph(
    'Проведена полноценная локальная приёмка Hermes-помощника по Element без подключения к Matrix. '
    'Проверка выполнена через CLI-запросы к локальному профилю helpdesk-агента.'
)

bullet_points = [
    'Сценарий: локальная dry-run приёмка без Matrix-подключения.',
    'Профиль: /root/element-hermes-assistant/docker-gateway/hermes-home',
    'Команда запуска: /root/.hermes/venv/bin/hermes chat -q ... -Q --max-turns 8',
    f"Всего тестовых вопросов: {len(results)}",
    f"Итог: PASS = {summary.get('PASS', 0)}, PARTIAL = {summary.get('PARTIAL', 0)}, FAIL = {summary.get('FAIL', 0)}",
]
for item in bullet_points:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2. Общий итог', level=1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Показатель'
hdr[1].text = 'PASS'
hdr[2].text = 'PARTIAL'
hdr[3].text = 'FAIL'
row = table.add_row().cells
row[0].text = 'Всего по всем вопросам'
row[1].text = str(summary.get('PASS', 0))
row[2].text = str(summary.get('PARTIAL', 0))
row[3].text = str(summary.get('FAIL', 0))

doc.add_paragraph('')
doc.add_heading('3. Сводка по разделам', level=1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Раздел'
hdr[1].text = 'PASS'
hdr[2].text = 'PARTIAL'
hdr[3].text = 'FAIL'
for sec_name, stats in sections.items():
    row = table.add_row().cells
    row[0].text = sec_name
    row[1].text = str(stats.get('PASS', 0))
    row[2].text = str(stats.get('PARTIAL', 0))
    row[3].text = str(stats.get('FAIL', 0))

doc.add_paragraph('')
doc.add_heading('4. Глобальные проверки качества', level=1)
for qc in quality_checks:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f"Q{qc['num']} — {qc['title']}: {qc['verdict']}")
    run.bold = True
    if qc.get('evidence'):
        doc.add_paragraph(qc['evidence'])

doc.add_heading('5. Выводы по готовности', level=1)
for item in [
    'Пользовательские ответы больше не содержат сырых шаблонных плейсхолдеров.',
    'Агент корректно обрабатывает темы логина, уведомлений, комнат, файлов, звонков, устройств и шифрования.',
    'На рискованных темах (ключи, Secure Backup, незнакомые сессии) агент предупреждает о последствиях и не предлагает опасные действия «наугад».',
    'По локальным правилам агент использует зафиксированные факты: login/password, без SSO, без VPN, BYOD разрешён, внешние участники запрещены без отдельного решения, эскалация — через бота.',
    'По результатам локальной приёмки профиль готов к следующему этапу — pre-production проверке перед подключением к Matrix.'
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6. Детальные результаты по вопросам', level=1)
for r in results:
    doc.add_heading(f"Q{r['num']}. {r['question']}", level=2)
    doc.add_paragraph(f"Раздел: {r['section']}")
    p = doc.add_paragraph()
    rr = p.add_run(f"Вердикт: {r['verdict']}")
    rr.bold = True
    doc.add_paragraph(f"Время ответа: {r['elapsed_sec']} сек")
    doc.add_paragraph('Ответ:')
    doc.add_paragraph(r['answer'])


doc.add_heading('7. Артефакты проверки', level=1)
for item in [
    'Markdown-отчёт: /root/element-hermes-assistant/docker-gateway/reports/acceptance_local_20260419_051026.md',
    'JSON-отчёт: /root/element-hermes-assistant/docker-gateway/reports/acceptance_local_20260419_051026.json',
    'Скрипт прогона: /root/element-hermes-assistant/docker-gateway/run_local_acceptance.py',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.save(out_path)
print(out_path)
